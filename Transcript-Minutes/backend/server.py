from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, Response
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import json
import re
from io import BytesIO
import tempfile

# AI Integration
from emergentintegrations.llm.chat import LlmChat, UserMessage
from dotenv import load_dotenv

# Document processing
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import PageTemplate, BaseDocTemplate
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import seaborn as sns
from datetime import datetime, timezone
import base64

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# AI Chat instance
API_KEY = os.environ.get('EMERGENT_LLM_KEY')
if not API_KEY:
    raise Exception("EMERGENT_LLM_KEY not found in environment variables")

# Define Models
class ActionItem(BaseModel):
    owner: str
    task: str
    deadline: Optional[str] = None
    status: str = "pending"
    confidence: float = 0.0

class Decision(BaseModel):
    decision: str
    context: str
    confidence: float = 0.0

class Topic(BaseModel):
    topic: str
    confidence: float = 0.0
    keywords: List[str] = []

class MeetingMinutes(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    summary: str
    agenda_items: List[str] = []
    action_items: List[ActionItem] = []
    decisions: List[Decision] = []
    participants: List[str] = []
    topics: List[Topic] = []
    original_transcript: str
    processed_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    processing_method: str = "ai"  # "ai" or "regex"

class TranscriptInput(BaseModel):
    transcript: str

# AI Processing Functions
async def process_transcript_with_ai(transcript: str) -> Dict[str, Any]:
    """Process transcript using AI (GPT-5)"""
    try:
        chat = LlmChat(
            api_key=API_KEY,
            session_id=f"transcript-{uuid.uuid4()}",
            system_message="""You are an expert meeting minutes processor. Extract structured information from meeting transcripts with high accuracy.
            
For each transcript, provide a JSON response with:
            1. summary: Clear bullet-point summary (2-4 bullets)
            2. agenda_items: List of topics/agenda items discussed
            3. action_items: Array of {"owner": "person", "task": "description", "deadline": "date or null", "status": "pending", "confidence": 0.0-1.0}
            4. decisions: Array of {"decision": "what was decided", "context": "background", "confidence": 0.0-1.0}
            5. participants: List of unique, valid participant names (exclude words like "date", "need", "time", etc.)
            6. topics: Array of {"topic": "discussion area", "confidence": 0.0-1.0, "keywords": ["key", "words"]}
            
Be accurate. If no action items exist, return empty array. Only include real participant names."""
        ).with_model("openai", "gpt-5")
        
        user_message = UserMessage(
            text=f"Process this meeting transcript and return structured JSON:\n\n{transcript}"
        )
        
        response = await chat.send_message(user_message)
        
        # Parse JSON response
        try:
            # Extract JSON from response if it contains additional text
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            if json_start != -1 and json_end != -1:
                json_str = response[json_start:json_end]
                result = json.loads(json_str)
            else:
                result = json.loads(response)
            
            # Validate and set default confidence scores
            for item in result.get('action_items', []):
                if 'confidence' not in item:
                    item['confidence'] = 0.8
            
            for item in result.get('decisions', []):
                if 'confidence' not in item:
                    item['confidence'] = 0.9
            
            for item in result.get('topics', []):
                if 'confidence' not in item:
                    item['confidence'] = 0.7
                if 'keywords' not in item:
                    item['keywords'] = []
            
            result['processing_method'] = 'ai'
            return result
            
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            print(f"Response: {response}")
            # Fall back to regex processing
            return process_transcript_with_regex(transcript)
            
    except Exception as e:
        print(f"AI processing error: {e}")
        # Fall back to regex processing
        return process_transcript_with_regex(transcript)

def process_transcript_with_regex(transcript: str) -> Dict[str, Any]:
    """Fallback regex processing for reliability"""
    
    # Summary - first few sentences or key points
    sentences = transcript.split('. ')
    summary_bullets = []
    for i, sentence in enumerate(sentences[:3]):
        if len(sentence.strip()) > 20:
            summary_bullets.append(f"• {sentence.strip()}")
    summary = "\n".join(summary_bullets) if summary_bullets else "Meeting summary from transcript"
    
    # Extract participants using common patterns
    participants = set()
    # Pattern 1: "John said" or "Mary mentioned"
    name_patterns = [
        r'([A-Z][a-z]+)\s+(?:said|mentioned|asked|replied|stated|noted)',
        r'([A-Z][a-z]+):\s',  # "John: we need to..."
        r'(?:from|by)\s+([A-Z][a-z]+)',
    ]
    
    for pattern in name_patterns:
        matches = re.findall(pattern, transcript)
        for match in matches:
            name = match.strip()
            if len(name) > 2 and name not in ['The', 'We', 'This', 'That', 'Need', 'Time', 'Date']:
                participants.add(name)
    
    # Action items - look for commitment patterns
    action_items = []
    action_patterns = [
        r'([A-Z][a-z]+)\s+(?:will|should|needs to|has to|must)\s+([^.]+)',
        r'(?:action|task|todo)(?:\s+for)?\s+([A-Z][a-z]+):\s*([^.]+)',
        r'([A-Z][a-z]+)\s+(?:agreed to|committed to)\s+([^.]+)'
    ]
    
    for pattern in action_patterns:
        matches = re.findall(pattern, transcript, re.IGNORECASE)
        for match in matches:
            if len(match) == 2:
                owner, task = match
                action_items.append({
                    "owner": owner,
                    "task": task.strip(),
                    "deadline": None,
                    "status": "pending",
                    "confidence": 0.6
                })
    
    # Decisions - look for decision patterns
    decisions = []
    decision_patterns = [
        r'(?:decided|agreed|determined)\s+(?:that|to)?\s*([^.]+)',
        r'(?:decision|conclusion):\s*([^.]+)',
    ]
    
    for pattern in decision_patterns:
        matches = re.findall(pattern, transcript, re.IGNORECASE)
        for match in matches:
            decisions.append({
                "decision": match.strip(),
                "context": "From meeting discussion",
                "confidence": 0.5
            })
    
    # Topics - simple keyword extraction
    topics = []
    topic_keywords = ['project', 'budget', 'timeline', 'deadline', 'team', 'client', 'proposal', 'meeting']
    for keyword in topic_keywords:
        if keyword.lower() in transcript.lower():
            topics.append({
                "topic": keyword.title(),
                "confidence": 0.4,
                "keywords": [keyword]
            })
    
    return {
        "summary": summary,
        "agenda_items": ["Meeting discussion", "Action items", "Next steps"],
        "action_items": action_items,
        "decisions": decisions,
        "participants": list(participants)[:10],  # Limit to 10
        "topics": topics,
        "processing_method": "regex"
    }

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(BytesIO(file_content))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing DOCX file: {str(e)}")

def create_simple_charts(minutes: MeetingMinutes) -> dict:
    """Generate simple charts for meeting insights"""
    charts = {}
    
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import numpy as np
        
        # 1. Action Items Distribution Chart
        if minutes.action_items and len(minutes.action_items) > 0:
            fig, ax = plt.subplots(figsize=(8, 6))
            owners = [item.owner for item in minutes.action_items]
            owner_counts = {}
            for owner in owners:
                owner_counts[owner] = owner_counts.get(owner, 0) + 1
            
            if owner_counts:
                ax.bar(list(owner_counts.keys()), list(owner_counts.values()))
                ax.set_title('Action Items by Owner')
                ax.set_xlabel('Owner')
                ax.set_ylabel('Count')
                
                plt.tight_layout()
                chart_path = '/tmp/action_items_chart.png'
                plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                charts['action_items'] = chart_path
                plt.close()
        
        # 2. Topics Confidence Chart
        if minutes.topics and len(minutes.topics) > 0:
            fig, ax = plt.subplots(figsize=(8, 6))
            topics = [topic.topic[:15] + '...' if len(topic.topic) > 15 else topic.topic for topic in minutes.topics]
            confidences = [topic.confidence * 100 for topic in minutes.topics]
            
            ax.barh(topics, confidences)
            ax.set_title('Topics Confidence')
            ax.set_xlabel('Confidence %')
            
            plt.tight_layout()
            chart_path = '/tmp/topics_chart.png'
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            charts['topics'] = chart_path
            plt.close()
            
    except Exception as e:
        print(f"Chart generation error: {e}")
    
    return charts

def generate_simple_pdf_report(minutes: MeetingMinutes) -> BytesIO:
    """Generate simplified PDF report"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=1*inch, rightMargin=1*inch,
                           topMargin=1*inch, bottomMargin=1*inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Title with date
    date_str = minutes.processed_at.strftime('%B %d, %Y')
    title = Paragraph(f'Meeting Minutes Report - {date_str}', styles['Title'])
    story.append(title)
    story.append(Spacer(1, 0.3*inch))
    
    # Meeting Info
    info_text = f"""
    <b>Processing Method:</b> {minutes.processing_method.upper()}<br/>
    <b>Participants:</b> {len(minutes.participants)}<br/>
    <b>Action Items:</b> {len(minutes.action_items)}<br/>
    <b>Decisions:</b> {len(minutes.decisions)}<br/>
    <b>Topics:</b> {len(minutes.topics)}
    """
    story.append(Paragraph(info_text, styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Summary
    story.append(Paragraph('Executive Summary', styles['Heading2']))
    summary_lines = minutes.summary.replace('\\\\n', '\n').replace('\\n', '\n').split('\n')
    for line in summary_lines:
        if line.strip():
            story.append(Paragraph(f'• {line.strip()}', styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Participants
    if minutes.participants:
        story.append(Paragraph('Participants', styles['Heading2']))
        participants_text = ', '.join(minutes.participants)
        story.append(Paragraph(participants_text, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    # Action Items
    if minutes.action_items:
        story.append(Paragraph('Action Items', styles['Heading2']))
        for item in minutes.action_items:
            action_text = f'<b>{item.owner}:</b> {item.task}'
            if item.deadline:
                action_text += f' <i>(Due: {item.deadline})</i>'
            action_text += f' <i>[{int(item.confidence * 100)}% confidence]</i>'
            story.append(Paragraph(action_text, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    # Decisions
    if minutes.decisions:
        story.append(Paragraph('Key Decisions', styles['Heading2']))
        for i, decision in enumerate(minutes.decisions, 1):
            story.append(Paragraph(f'<b>{i}.</b> {decision.decision}', styles['Normal']))
            story.append(Paragraph(f'<i>Context: {decision.context}</i>', styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    # Topics
    if minutes.topics:
        story.append(Paragraph('Discussion Topics', styles['Heading2']))
        for topic in minutes.topics:
            topic_text = f'<b>{topic.topic}</b> ({int(topic.confidence * 100)}% confidence)'
            if topic.keywords:
                topic_text += f' - Keywords: {", ".join(topic.keywords[:3])}'
            story.append(Paragraph(topic_text, styles['Normal']))
            story.append(Spacer(1, 0.05*inch))
    
    # Add simple charts if available
    charts = create_simple_charts(minutes)
    if charts:
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph('Meeting Analytics', styles['Heading2']))
        
        for chart_name, chart_path in charts.items():
            try:
                img = Image(chart_path, width=5*inch, height=3*inch)
                story.append(img)
                story.append(Spacer(1, 0.2*inch))
            except Exception as e:
                print(f"Error adding chart {chart_name}: {e}")
    
    # Footer info
    story.append(Spacer(1, 0.5*inch))
    footer_text = f'Generated on {minutes.processed_at.strftime("%B %d, %Y at %I:%M %p")} by Meeting Minutes AI'
    story.append(Paragraph(footer_text, styles['Normal']))
    
    doc.build(story)
    
    # Clean up chart files
    for chart_path in charts.values():
        try:
            import os
            os.remove(chart_path)
        except:
            pass
    
    buffer.seek(0)
    return buffer

def generate_simple_docx_report(minutes: MeetingMinutes) -> BytesIO:
    """Generate simplified DOCX report"""
    buffer = BytesIO()
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Meeting Minutes Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Generated on {minutes.processed_at.strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meeting Info
    doc.add_heading('Meeting Information', level=1)
    info_para = doc.add_paragraph()
    info_para.add_run(f'Processing Method: ').bold = True
    info_para.add_run(f'{minutes.processing_method.upper()}\n')
    info_para.add_run(f'Participants: ').bold = True
    info_para.add_run(f'{len(minutes.participants)}\n')
    info_para.add_run(f'Action Items: ').bold = True
    info_para.add_run(f'{len(minutes.action_items)}\n')
    info_para.add_run(f'Decisions: ').bold = True
    info_para.add_run(f'{len(minutes.decisions)}\n')
    info_para.add_run(f'Topics: ').bold = True
    info_para.add_run(f'{len(minutes.topics)}')
    
    # Summary
    doc.add_heading('Executive Summary', level=1)
    summary_lines = minutes.summary.replace('\\\\n', '\n').replace('\\n', '\n').split('\n')
    for line in summary_lines:
        if line.strip():
            p = doc.add_paragraph(line.strip(), style='List Bullet')
    
    # Participants
    if minutes.participants:
        doc.add_heading('Participants', level=1)
        participants_para = doc.add_paragraph(', '.join(minutes.participants))
    
    # Action Items
    if minutes.action_items:
        doc.add_heading('Action Items', level=1)
        for item in minutes.action_items:
            para = doc.add_paragraph()
            para.add_run(f'{item.owner}: ').bold = True
            para.add_run(item.task)
            if item.deadline:
                para.add_run(f' (Due: {item.deadline})').italic = True
            para.add_run(f' [{int(item.confidence * 100)}% confidence]').italic = True
    
    # Decisions
    if minutes.decisions:
        doc.add_heading('Key Decisions', level=1)
        for i, decision in enumerate(minutes.decisions, 1):
            para = doc.add_paragraph()
            para.add_run(f'{i}. ').bold = True
            para.add_run(decision.decision)
            
            context_para = doc.add_paragraph()
            context_para.add_run('Context: ').italic = True
            context_para.add_run(decision.context).italic = True
    
    # Topics
    if minutes.topics:
        doc.add_heading('Discussion Topics', level=1)
        for topic in minutes.topics:
            para = doc.add_paragraph()
            para.add_run(topic.topic).bold = True
            para.add_run(f' ({int(topic.confidence * 100)}% confidence)')
            if topic.keywords:
                para.add_run(f' - Keywords: {", ".join(topic.keywords[:3])}').italic = True
    
    # Add charts
    charts = create_simple_charts(minutes)
    if charts:
        doc.add_heading('Meeting Analytics', level=1)
        for chart_name, chart_path in charts.items():
            try:
                doc.add_picture(chart_path, width=Inches(6))
                doc.add_paragraph()
            except Exception as e:
                print(f"Error adding chart to DOCX: {e}")
    
    # Clean up chart files
    for chart_path in charts.values():
        try:
            import os
            os.remove(chart_path)
        except:
            pass
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# API Routes
@api_router.post("/process-transcript", response_model=MeetingMinutes)
async def process_transcript(input: TranscriptInput):
    """Process a meeting transcript (text input)"""
    if not input.transcript or len(input.transcript.strip()) < 10:
        raise HTTPException(status_code=400, detail="Transcript must be at least 10 characters long")
    
    try:
        # Process with AI
        result = await process_transcript_with_ai(input.transcript)
        
        # Create MeetingMinutes object
        summary_text = result.get('summary', 'Meeting summary')
        if isinstance(summary_text, list):
            summary_text = "\n".join(summary_text)
        # Convert escaped newlines to actual newlines
        summary_text = summary_text.replace('\\n', '\n')
        
        minutes = MeetingMinutes(
            summary=summary_text,
            agenda_items=result.get('agenda_items', []),
            action_items=[ActionItem(**item) for item in result.get('action_items', [])],
            decisions=[Decision(**decision) for decision in result.get('decisions', [])],
            participants=result.get('participants', []),
            topics=[Topic(**topic) for topic in result.get('topics', [])],
            original_transcript=input.transcript,
            processing_method=result.get('processing_method', 'ai')
        )
        
        return minutes
        
    except Exception as e:
        print(f"Error processing transcript: {e}")
        raise HTTPException(status_code=500, detail="Error processing transcript")

@api_router.post("/upload-transcript", response_model=MeetingMinutes)
async def upload_transcript(file: UploadFile = File(...)):
    """Upload and process a transcript file (DOCX or TXT)"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    # Check file type
    allowed_types = ['.txt', '.docx']
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in allowed_types:
        raise HTTPException(status_code=400, detail="Only .txt and .docx files are supported")
    
    try:
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        if file_ext == '.docx':
            transcript_text = extract_text_from_docx(content)
        else:  # .txt
            transcript_text = content.decode('utf-8')
        
        if len(transcript_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Transcript must be at least 10 characters long")
        
        # Process with AI
        result = await process_transcript_with_ai(transcript_text)
        
        # Create MeetingMinutes object
        summary_text = result.get('summary', 'Meeting summary')
        if isinstance(summary_text, list):
            summary_text = "\n".join(summary_text)
        # Convert escaped newlines to actual newlines
        summary_text = summary_text.replace('\\n', '\n')
        
        minutes = MeetingMinutes(
            summary=summary_text,
            agenda_items=result.get('agenda_items', []),
            action_items=[ActionItem(**item) for item in result.get('action_items', [])],
            decisions=[Decision(**decision) for decision in result.get('decisions', [])],
            participants=result.get('participants', []),
            topics=[Topic(**topic) for topic in result.get('topics', [])],
            original_transcript=transcript_text,
            processing_method=result.get('processing_method', 'ai')
        )
        
        return minutes
        
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="Unable to read file. Please ensure it's a valid text file.")
    except Exception as e:
        print(f"Error processing uploaded file: {e}")
        raise HTTPException(status_code=500, detail="Error processing uploaded file")

@api_router.post("/export-pdf")
async def export_pdf(minutes_data: MeetingMinutes):
    """Export meeting minutes as PDF with charts"""
    try:
        pdf_buffer = generate_simple_pdf_report(minutes_data)
        
        return Response(
            content=pdf_buffer.getvalue(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=meeting-minutes-{minutes_data.processed_at.strftime('%Y%m%d-%H%M')}.pdf"
            }
        )
    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

@api_router.post("/export-docx")
async def export_docx(minutes_data: MeetingMinutes):
    """Export meeting minutes as DOCX with charts"""
    try:
        docx_buffer = generate_simple_docx_report(minutes_data)
        
        return Response(
            content=docx_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=meeting-minutes-{minutes_data.processed_at.strftime('%Y%m%d-%H%M')}.docx"
            }
        )
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

@api_router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "ai_integration": "enabled"}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
